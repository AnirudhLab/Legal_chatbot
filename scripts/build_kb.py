import os
import pdfplumber
import docx
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document

# ------------ CONFIG ------------
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

RAW_DOCS_PATH = "../data/raw_docs"
INDEX_PATH = "../index/openai_index"
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50
EMBED_MODEL = "text-embedding-3-small"
MAX_CHARS = 7500
# --------------------------------


def extract_text(file_path):
    try:
        if file_path.endswith(".pdf"):
            with pdfplumber.open(file_path) as pdf:
                return "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()])
        elif file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        elif file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
    except Exception as e:
        print(f"[⚠️] Skipping {file_path}: {e}")
    return ""


def load_and_chunk_documents(folder_path):
    all_chunks = []
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

    for filename in os.listdir(folder_path):
        path = os.path.join(folder_path, filename)
        text = extract_text(path)
        if text:
            chunks = splitter.split_text(text)
            for chunk in chunks:
                cleaned = chunk.strip()
                if cleaned and len(cleaned) <= MAX_CHARS:
                    all_chunks.append(Document(page_content=cleaned, metadata={"source": filename}))
                elif len(cleaned) > MAX_CHARS:
                    print(f"[⚠️] Skipping overlong chunk in {filename} ({len(cleaned)} chars)")
    return all_chunks


def create_faiss_index(docs):
    print("[🔑] Using OpenAI API for embeddings...")

    embeddings = OpenAIEmbeddings(
        model=EMBED_MODEL,
        api_key=OPENAI_API_KEY
    )

    texts = []
    metadatas = []

    for idx, doc in enumerate(docs):
        if isinstance(doc.page_content, str) and doc.page_content.strip():
            texts.append(doc.page_content.strip())
            metadatas.append(doc.metadata)
        else:
            print(f"[⚠️] Invalid or empty doc at index {idx}")

    print(f"[🧪] Prepared {len(texts)} clean chunks for embedding.")

    print("[💾] Creating FAISS vector store...")
    vectorstore = FAISS.from_texts(texts=texts, embedding=embeddings, metadatas=metadatas)
    vectorstore.save_local(INDEX_PATH)
    print(f"[✅] FAISS vector store saved to: {INDEX_PATH}")


def main():
    print("[📄] Loading and chunking documents...")
    docs = load_and_chunk_documents(RAW_DOCS_PATH)

    print(f"[📦] Total valid chunks: {len(docs)}")
    create_faiss_index(docs)


if __name__ == "__main__":
    main()
